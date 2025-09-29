import os
import magic
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import aiofiles
import asyncio
from sentence_transformers import SentenceTransformer
import numpy as np
import json
from backend.api.models.database import Document, DocumentChunk
from sqlalchemy.orm import Session

class DocumentProcessor:
    def __init__(self):
        self.model = None
        self.chunk_sizes = {
            'pdf': 1000,
            'docx': 800,
            'txt': 1200,
            'csv': 500
        }
        
    async def load_model(self):
        """Load the embedding model"""
        if self.model is None:
            self.model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
        return self.model

    async def process_documents(self, file_paths: List[str], db_session: Session) -> Dict[str, Any]:
        """Process multiple documents"""
        results = {
            'processed': 0,
            'failed': 0,
            'documents': []
        }
        
        await self.load_model()
        
        for file_path in file_paths:
            try:
                document_info = await self.process_single_document(file_path, db_session)
                results['documents'].append(document_info)
                results['processed'] += 1
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                results['failed'] += 1
                results['documents'].append({
                    'filename': os.path.basename(file_path),
                    'status': 'error',
                    'error': str(e)
                })
        
        return results

    async def process_single_document(self, file_path: str, db_session: Session) -> Dict[str, Any]:
        """Process a single document with better error handling"""
        filename = os.path.basename(file_path)
        file_type = self._detect_file_type(file_path)
        
        print(f"Processing document: {filename} (type: {file_type})")
        
        # Create document record
        document = Document(
            filename=filename,
            file_type=file_type,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            processed=1  # Processing
        )
        db_session.add(document)
        db_session.commit()
        
        try:
            # Extract content
            content = await self._extract_content(file_path, file_type)
            print(f"Extracted {len(content)} characters from {filename}")
            
            if not content or len(content.strip()) == 0:
                raise Exception("No content extracted from document")
            
            # Update document with content
            document.content = content
            document.processed = 2  # Completed
            db_session.commit()
            
            # Chunk and embed
            chunks = self.dynamic_chunking(content, file_type)
            print(f"Created {len(chunks)} chunks from {filename}")
            
            if chunks:
                embeddings = await self._generate_embeddings_batch(chunks)
                print(f"Generated {len(embeddings)} embeddings for {filename}")
                
                # Store chunks
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    doc_chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk,
                        embedding=json.dumps(embedding.tolist()),
                        tokens_count=len(chunk.split())
                    )
                    db_session.add(doc_chunk)
                
                db_session.commit()
            else:
                print(f"No chunks created for {filename}")
            
            return {
                'id': document.id,
                'filename': filename,
                'file_type': file_type,
                'status': 'completed',
                'chunks_count': len(chunks),
                'content_length': len(content)
            }
            
        except Exception as e:
            print(f"Error processing {filename}: {e}")
            document.processed = 3  # Error
            db_session.commit()
            raise e

    def dynamic_chunking(self, content: str, doc_type: str) -> List[str]:
        """Intelligent chunking based on document type and structure"""
        chunk_size = self.chunk_sizes.get(doc_type, 512)
        
        if doc_type == 'pdf':
            return self._chunk_pdf_content(content, chunk_size)
        elif doc_type == 'docx':
            return self._chunk_docx_content(content, chunk_size)
        elif doc_type == 'csv':
            return self._chunk_csv_content(content, chunk_size)
        else:
            return self._chunk_generic_content(content, chunk_size)

    def _chunk_pdf_content(self, content: str, chunk_size: int) -> List[str]:
        """Chunk PDF content preserving sections"""
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = ""
        
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += " " + paragraph
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

    def _chunk_docx_content(self, content: str, chunk_size: int) -> List[str]:
        """Chunk DOCX content preserving paragraphs"""
        return self._chunk_generic_content(content, chunk_size)

    def _chunk_csv_content(self, content: str, chunk_size: int) -> List[str]:
        """Chunk CSV content row by row"""
        lines = content.split('\n')
        chunks = []
        current_chunk = ""
        
        for line in lines:
            if len(current_chunk) + len(line) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = line
            else:
                current_chunk += "\n" + line
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks

    def _chunk_generic_content(self, content: str, chunk_size: int) -> List[str]:
        """Generic chunking for text content"""
        words = content.split()
        chunks = []
        current_chunk = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) > chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                current_chunk = [word]
                current_length = len(word)
            else:
                current_chunk.append(word)
                current_length += len(word) + 1
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extract content from different file types"""
        if file_type == 'pdf':
            return self._extract_pdf_content(file_path)
        elif file_type == 'docx':
            return self._extract_docx_content(file_path)
        elif file_type == 'csv':
            return self._extract_csv_content(file_path)
        elif file_type == 'txt':
            return await self._extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF"""
        content = ""
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        return content

    def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        content = ""
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        return content

    def _extract_csv_content(self, file_path: str) -> str:
        """Extract content from CSV"""
        df = pd.read_csv(file_path)
        return df.to_string()

    async def _extract_txt_content(self, file_path: str) -> str:
        """Extract text from TXT"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()

    def _detect_file_type(self, file_path: str) -> str:
        """Detect file type using python-magic"""
        mime = magic.Magic(mime=True)
        mime_type = mime.from_file(file_path)
        
        if 'pdf' in mime_type:
            return 'pdf'
        elif 'word' in mime_type or 'officedocument' in mime_type:
            return 'docx'
        elif 'csv' in mime_type or 'spreadsheet' in mime_type:
            return 'csv'
        elif 'text' in mime_type:
            return 'txt'
        else:
            # Fallback to extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.pdf']:
                return 'pdf'
            elif ext in ['.docx', '.doc']:
                return 'docx'
            elif ext in ['.csv', '.xlsx', '.xls']:
                return 'csv'
            elif ext in ['.txt']:
                return 'txt'
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")

    async def _generate_embeddings_batch(self, chunks: List[str]) -> List[np.ndarray]:
        """Generate embeddings in batches"""
        if not self.model:
            await self.load_model()
        
        embeddings = self.model.encode(chunks, batch_size=32, show_progress_bar=False)
        return embeddings